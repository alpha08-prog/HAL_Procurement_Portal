import re, json
from pathlib import Path
from docx import Document

def _read(path):
    doc = Document(path)
    paras = [p.text for p in doc.paragraphs if p.text.strip()]
    tables = [[[c.text.strip() for c in r.cells] for r in t.rows] for t in doc.tables]
    return paras, tables

def _flatten(paras, tables):
    lines = list(paras)
    for tbl in tables:
        for row in tbl:
            if any(row):
                lines.append(" | ".join(row))
    return "\n".join(lines)

def _dedupe(seq):
    seen, out = set(), []
    for v in seq:
        v = v.strip()
        if v and v not in seen:
            seen.add(v); out.append(v)
    return out

def _match(text, spec):
    if isinstance(spec, str):
        spec = {"pattern": spec}
    pat = spec["pattern"]
    grp = spec.get("group", 1)
    multi = spec.get("multi", False)
    flags = spec.get("flags", re.IGNORECASE | re.DOTALL)
    if multi:
        raw = re.findall(pat, text, flags)
        if raw and isinstance(raw[0], tuple):
            return _dedupe([m[grp - 1] for m in raw])
        return _dedupe(raw)
    m = re.search(pat, text, flags)
    if not m:
        return None
    try:
        return m.group(grp).strip()
    except IndexError:
        return m.group(0).strip()

def extract_tables(path):
    return _read(path)[1]

def read_text(path):
    paras, tables = _read(path)
    return _flatten(paras, tables)

def extract_from_docx(path, patterns):
    paras, tables = _read(path)
    text = _flatten(paras, tables)
    out = {"_source": Path(path).name, "_raw_char_count": len(text), "tables": tables}
    for field, spec in patterns.items():
        out[field] = _match(text, spec)
    return out

def extract_from_docx_to_json(path, patterns, output_path):
    data = extract_from_docx(path, patterns)
    p = Path(output_path)
    p.parent.mkdir(parents=True, exist_ok=True)
    p.write_text(json.dumps(data, indent=2, ensure_ascii=False), encoding="utf-8")
    return data
